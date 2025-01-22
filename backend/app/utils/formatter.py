import logging
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# Key APA Guidelines for Theses
#
# 1. Title Page:
#   - Title: Centered, bold, and written in title case.
#   - Author Name: Centered, with the first name followed by the last name.
#   - Institution: Full name of the institution.
#   - Course, instructor, and submission date (if applicable).
#
# 2. Abstract:
#   - Title: Centered and bold.
#   - Content: A single paragraph, double-spaced, and no more than 250 words.
#
# 3. Body Text:
#   - Use a consistent font (e.g., Times New Roman, 12pt).
#   - Double-space throughout, with no extra spaces between paragraphs.
#   - Margins: 1 inch on all sides.
#
# 4. References:
#   - Use a hanging indent (first line flush left, subsequent lines indented).
#   - Format references according to APA rules:
#     - Author last name, first initial.
#     - Year in parentheses.
#     - Italicize book titles or journal names.
#     - Include DOIs or URLs where applicable.
#
# 5. Headings:
#   - Follow APA heading levels (5 levels, each with a specific formatting style).


def format_to_apa(thesis_data, references):
    """
    Formats thesis data into APA-compliant Word format.
    :param thesis_data: Dictionary containing thesis metadata (title, author, abstract).
    :param references: List of dictionaries containing reference metadata.
    :return: Path to the formatted APA Word document.
    """
    # Validate thesis data
    try:
        validate_thesis_data(thesis_data)
    except ValueError as e:
        logger.error(f"Invalid thesis data: {e}")
        raise

    # Validate references
    try:
        validate_references(references)
    except ValueError as e:
        logger.error(f"Invalid reference data: {e}")
        raise

    try:
        # Create a new Word document
        doc = Document()

        # Add Title Page
        add_title_page(doc, thesis_data)

        # Add Abstract Section
        if thesis_data.get("abstract"):
            add_abstract_section(doc, thesis_data["abstract"])

        # Add Content Section
        if thesis_data.get("content"):
            add_main_body(doc, thesis_data["title"], thesis_data["content"])

        # Add References Section
        if references:
            add_references_section(doc, references)

        # Apply double-spacing
        apply_double_spacing(doc)

        # Save document to a temporary location
        temp_file = "/tmp/apa_formatted_thesis.docx"
        if os.path.exists(temp_file):
            os.remove(temp_file)

        doc.save(temp_file)

        logger.info(f"APA-formatted document saved at {temp_file}")
        return temp_file

    except Exception as e:
        logger.error(f"An error occurred while formatting the APA document: {e}")
        raise RuntimeError("Failed to format the APA document.") from e


def validate_thesis_data(thesis_data):
    """
    Validates the thesis metadata.
    :param thesis_data: Dictionary containing thesis metadata.
    """
    required_fields = [
        "title",
        "student",
        "institution",
        "instructor",
        "course",
        "abstract",
        "submission_date",
    ]
    missing_fields = [field for field in required_fields if not thesis_data.get(field)]
    if missing_fields:
        raise ValueError(
            f"Missing required thesis data fields: {', '.join(missing_fields)}"
        )


def validate_references(references):
    """
    Validates the reference metadata.
    :param references: List of dictionaries containing reference metadata.
    """
    if not isinstance(references, list):
        raise ValueError("References should be a list of dictionaries.")
    for i, ref in enumerate(references):
        if not isinstance(ref, dict):
            raise ValueError(f"Reference at index {i} is not a dictionary.")
        required_fields = ["author", "title", "publication_year"]
        for field in required_fields:
            if field not in ref:
                raise ValueError(
                    f"Reference {i} is missing the required field: {field}."
                )
        if not isinstance(ref["publication_year"], int):
            raise ValueError(
                f"Reference {i}: 'publication_year' must be an integer. Got {type(ref['publication_year']).__name__}."
            )


def add_title_page(doc, thesis_data):
    """
    Adds the title page as per APA guidelines.
    """
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = f"Running head: {thesis_data['title'].upper()[:50]}"
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add the title, bold and centered
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(thesis_data["title"])
    title_run.bold = True
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add the rest of the title page information, centered
    doc.add_paragraph(f"{thesis_data['student']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"{thesis_data['institution']}").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph(f"{thesis_data['course']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Instructor: {thesis_data['instructor']}").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph(
        f"Submission Date: {thesis_data['submission_date']}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_abstract_section(doc, abstract):
    """
    Adds the abstract section as per APA guidelines.
    """
    doc.add_page_break()
    doc.add_heading("Abstract", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    abstract_paragraph = doc.add_paragraph(abstract)
    abstract_paragraph.paragraph_format.first_line_indent = None
    abstract_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_heading(doc, text, level):
    """
    Adds a heading with APA formatting.
    :param doc: Word document object.
    :param text: Heading text.
    :param level: APA heading level (1-5).
    """
    heading = doc.add_heading(level=level)
    heading.text = text
    heading.bold = True
    if level == 1:
        heading.alignment = 1  # Centered
    else:
        heading.alignment = 0  # Left-aligned


def add_main_body(doc, title, content):
    """
    Adds the main body as per APA guidelines.
    """
    doc.add_page_break()
    doc.add_paragraph(title).bold = True
    doc.add_paragraph(content).alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.paragraphs[-1].paragraph_format.first_line_indent = 720  # 0.5 inch indent


def add_references_section(doc, references):
    """
    Adds the references section to the document as per APA guidelines.
    :param doc: Word document object.
    :param references: List of dictionaries containing reference metadata.
    """
    doc.add_page_break()
    doc.add_heading("References", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, ref in enumerate(references):
        try:
            # Build citation string
            citation = f"{ref['author']} ({ref['publication_year']}). {ref['title']}."
            if ref.get("journal"):
                citation += f" {ref['journal']}."
            if ref.get("doi"):
                citation += f" DOI: {ref['doi']}"

            # Add reference paragraph
            ref_paragraph = doc.add_paragraph(citation)
            ref_paragraph.style = doc.styles["Normal"]

            # Apply hanging indent (720 twips = 0.5 inch)
            if ref_paragraph.paragraph_format.left_indent is None:
                ref_paragraph.paragraph_format.left_indent = 0
            ref_paragraph.paragraph_format.left_indent += 720  # Hanging indent in twips

        except KeyError as e:
            logger.error(f"Missing key in reference {i}: {e}. Reference: {ref}")
            raise ValueError(f"Reference {i} is missing required fields.") from e
        except Exception as e:
            logger.error(f"Error adding reference {i}: {e}. Reference: {ref}")
            raise RuntimeError("Failed to add references section.") from e


def apply_global_styles(doc):
    """
    Applies global APA styles (double-spacing, margins) to the document.
    :param doc: Word document object.
    """
    # Double-spacing
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 2.0

    # Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = 1
        section.bottom_margin = 1
        section.left_margin = 1
        section.right_margin = 1


def apply_double_spacing(doc):
    """
    Applies double-spacing to all paragraphs in the document.
    :param doc: Word document object.
    """
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 2.0
