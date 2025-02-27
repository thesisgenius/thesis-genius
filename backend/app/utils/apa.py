# apa.py (Illustrative Example)

import datetime
import io
import re
from typing import Dict

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def create_apa_docx(data):
    """
    Generates a .docx that meets the 'Diss_Format_Checklist' and
    looks like 'Diss_Template' as closely as possible.

    'data' is expected to have keys for:
      data["cover"]             => {title, author, institution, etc.}
      data["copyright"]         => {year, text?}
      data["signature"]         => {defense_date, chair_name, ...}
      data["abstract"]          => { title, body }
      data["dedication"]        => { text } (optional)
      data["acknowledgments"]   => { text } (optional)
      data["toc"]               => [ ... ]  # or a single string for a placeholder
      data["list_of_tables"]    => [ ... ]
      data["list_of_figures"]   => [ ... ]
      data["chapters"]          => [ {heading, content}, ... ]
      data["references"]        => [ {some reference fields}, ... ]
      data["footnotes"]         => [ "Footnote1", "Footnote2", ...]
      data["appendices"]        => [ {title, content}, ...]
    or adapt to your aggregator shape.
    """

    doc = Document()

    # 1) Set margins 1", TNR 12pt, double spaced, left aligned (ragged right).
    _set_global_format(doc)

    # (A) Title Page (no page number)
    # We'll put it in doc.sections[0], which has no page number or roman.
    _make_title_page(doc, data.get("cover", {}))

    # Now we create a new section for the rest of the front matter,
    # which uses Roman numerals ii, iii, iv, etc., bottom center.
    front_matter_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _set_roman_pagination(front_matter_sec, start=2)

    # (B) Copyright Page => "ii"
    _make_copyright_page(doc, data.get("copyright", {}))

    # (C) Signature Page => "iii"
    _make_signature_page(doc, data.get("signature", {}))

    # (D) Abstract => "iv", with up to 350 words,
    # plus "title" if you want to reaffirm
    _make_abstract_page(doc, data.get("abstract", {}))

    # (E) Dedication (optional)
    if data.get("dedication", {}).get("text"):
        _make_dedication_page(doc, data["dedication"]["text"])

    # (F) Acknowledgments (optional)
    if data.get("acknowledgments", {}).get("text"):
        _make_acknowledgments_page(doc, data["acknowledgments"]["text"])

    # (G) Table of Contents
    _make_table_of_contents_page(doc, data.get("toc", []))

    # (H) List of Tables
    _make_list_of_tables_page(doc, data.get("list_of_tables", []))

    # (I) List of Figures
    _make_list_of_figures_page(doc, data.get("list_of_figures", []))

    # Now we start the MAIN TEXT => new section with Arabic pagination from 1
    main_text_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _set_arabic_pagination(main_text_sec, start=1)

    # (J) Chapters (introduction, method, etc.)
    # Each chapter can start on a new page or not, as you prefer.
    for ch in data.get("chapters", []):
        heading = ch.get("name", "Untitled Chapter")
        content = ch.get("content", "No content provided.")
        doc.add_heading(heading, level=1)
        doc.add_paragraph(_strip_html(content))
        doc.add_page_break()

    # (K) References
    doc.add_heading("References", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    refs = data.get("references", [])
    if refs:
        for r in refs:
            # naive approach
            doc.add_paragraph(_strip_html(str(r)), style="References")
    else:
        doc.add_paragraph("No references provided.")
    doc.add_page_break()

    # (L) Footnotes (if not embedded at page bottom)
    foots = data.get("footnotes", [])
    doc.add_heading("Footnotes", level=1)
    if foots:
        for i, foot in enumerate(foots, 1):
            doc.add_paragraph(f"{i}. {_strip_html(foot)}")
    else:
        doc.add_paragraph("No footnotes provided.")
    doc.add_page_break()

    # (M) Tables
    tabs = data.get("tables", [])
    doc.add_heading("Tables", level=1)
    if tabs:
        for t in tabs:
            cap = t.get("caption", "Untitled Table")
            doc.add_paragraph(_strip_html(cap))
            # possibly doc.add_picture(...)
    else:
        doc.add_paragraph("No tables provided.")
    doc.add_page_break()

    # (N) Figures
    figs = data.get("figures", [])
    doc.add_heading("Figures", level=1)
    if figs:
        for f in figs:
            cap = f.get("caption", "Untitled Figure")
            doc.add_paragraph(_strip_html(cap))
            # doc.add_picture(...) if you want
    else:
        doc.add_paragraph("No figures provided.")
    doc.add_page_break()

    # (O) Appendices
    apps = data.get("appendices", [])
    if not apps:
        # Possibly have a placeholder
        doc.add_heading("Appendices", level=1)
        doc.add_paragraph("No appendices provided.")
    else:
        for app in apps:
            ap_title = app.get("title", "Untitled Appendix")
            ap_content = app.get("content", "")
            doc.add_heading(f"Appendix: {ap_title}", level=1)
            doc.add_paragraph(_strip_html(ap_content))
            doc.add_page_break()

    # done
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


# -------------------- HELPER METHODS --------------------


def _set_global_format(doc: Document):
    """Apply 1-inch margins, TNR 12pt, double spacing, left alignment (ragged right)."""
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)
        # Typically, the checklist wants page # bottom center: we'll handle below.

    # Set Normal style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # ragged right
    style.paragraph_format.left_indent = Inches(0.5)  # .5" paragraph indent


def _set_roman_pagination(section, start=2):
    """
    Use Roman numerals (ii, iii, iv, ...) at the bottom center.
    'start=2' means the second page is 'ii', third is 'iii', etc.
    """
    # page numbering at bottom center => override default
    sect_footer = section.footer
    par = sect_footer.paragraphs[0]
    par.text = ""  # clear
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run()
    _add_page_field(run, num_format="roman", start=start)


def _set_arabic_pagination(section, start=1):
    """Arabic page numbering (1,2,3...) at bottom center, starting from 'start'."""
    sect_footer = section.footer
    par = sect_footer.paragraphs[0]
    par.text = ""
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run()
    _add_page_field(run, num_format="arabic", start=start)


def _add_page_field(run, num_format="roman", start=1):
    """
    Insert a page field with a specific format (roman or arabic)
    and start number, e.g. <w:pgNumType w:fmt="roman" w:start="2" />
    python-docx doesn't have a direct convenience method, so we hack the XML.
    """
    # For the actual 'section' object, we can set .start_type, .page_start, etc.
    # but that doesn't easily let us do roman vs. arabic mid-doc.
    # We'll forcibly insert a <w:fldSimple> with a PAGE field.
    # The format (roman or decimal) is handled in <w:pgNumType> for that section if we want to set it.

    # We can also do something like:
    # section.start_type = WD_SECTION_START.CONTINUOUS
    # But to set the numbering format or start in python-docx is tricky.
    # We'll attempt a simpler approach, but you might need additional section property hacks.

    # minimal approach: just insert a PAGE field => user must update field in Word for correct numbering
    # Or we create a separate method to set the docx sectPr w:pgNumType. We'll do that below if you want.

    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    # you can try adding a format switch: 'PAGE \* ROMAN' but that doesn't always hold
    # e.g.: 'PAGE \\* ROMAN'
    # Then rely on Word to interpret it.
    # There's no guaranteed python-docx for w:pgNumType with different starts easily.

    ftext = OxmlElement("w:t")
    ftext.text = "1"
    fld.append(ftext)
    run._r.append(fld)

    # If you want to forcibly set the actual numbering in docx section,
    # you'd do something like:
    sectPr = run._r.getparent().getparent().getparent()  # <w:sectPr>
    if sectPr is not None and sectPr.tag.endswith("sectPr"):
        # create or find w:pgNumType
        pgNumType = None
        for elem in sectPr:
            if elem.tag.endswith("pgNumType"):
                pgNumType = elem
                break
        if pgNumType is None:
            pgNumType = OxmlElement("w:pgNumType")
            sectPr.append(pgNumType)

        # set format
        if num_format == "roman":
            pgNumType.set(qn("w:fmt"), "roman")
        else:
            pgNumType.set(qn("w:fmt"), "decimal")

        # set start
        pgNumType.set(qn("w:start"), str(start))


def _make_title_page(doc: Document, cover: Dict[str, str]):
    """No page number, no running head. Just big bold title in center, plus author info."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Insert blank lines to push down
    for _ in range(8):
        doc.add_paragraph()

    title = cover.get("title", "UNTITLED DISSERTATION")
    p_title = doc.add_paragraph(title)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.runs[0]
    r_title.font.size = Pt(12)
    r_title.bold = True

    # Next lines: your doc template might want "by Jane Q. Student" or "in partial fulfillment..."
    author = cover.get("author", "Firstname Lastname")
    institution = cover.get("affiliation", "University of Whatever")
    course = cover.get("course", "Course Title")
    instructor = cover.get("instructor", "Instructor Name")
    date_str = cover.get("due_date", datetime.date.today().strftime("%Y-%m-%d"))

    doc.add_paragraph(author).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(institution).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(course).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(instructor).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(date_str).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Possibly more text, e.g.
    # "In Partial Fulfillment of the Requirements for the Degree of Doctor of Psychology"
    doc.add_page_break()


def _make_copyright_page(doc, cdata):
    """Roman numeral ii at bottom center. Typically: '© Year by Name. All Rights Reserved.'"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # new page is already from the next section
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    year = cdata.get("year", datetime.date.today().year)
    name = cdata.get("name", "Firstname Lastname")
    text = f"© {year} by {name}\nAll Rights Reserved"
    paragraph.add_run(text)
    doc.add_page_break()


def _make_signature_page(doc, sdata):
    """
    Roman numeral iii.
    Typically:
      'This dissertation by X has been approved by the committee members...
       Chair's name, date, etc.'
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Signature Page").bold = True

    doc.add_paragraph(
        "\nThis dissertation by "
        + sdata.get("name", "Firstname Lastname")
        + " has been approved by the committee members..."
    )
    doc.add_paragraph("_____________________________________________________")
    doc.add_paragraph("Chair: " + sdata.get("chair", "Chair Name"))
    doc.add_paragraph("_____________________________________________________")
    doc.add_paragraph(
        "Date of Dissertation Defense: "
        + sdata.get("due_date", datetime.date.today().strftime("%Y-%m-%d"))
    )
    doc.add_page_break()


def _make_abstract_page(doc, adata):
    """Roman numeral iv. Usually 350-word limit, plus KEYWORDS if needed."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    abstract_title = adata.get("title", "Abstract")
    body = _strip_html(adata.get("body", "No abstract provided."))
    doc.add_heading(abstract_title, level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(body)
    # optional keywords
    if "keywords" in adata:
        keys = adata["keywords"]  # e.g. a list of strings
        doc.add_paragraph("KEYWORDS: " + ", ".join(f'"{k}"' for k in sorted(keys)))
    doc.add_page_break()


def _make_dedication_page(doc, text):
    doc.add_heading("Dedication", level=1)
    doc.add_paragraph(_strip_html(text))
    doc.add_page_break()


def _make_acknowledgments_page(doc, text):
    doc.add_heading("Acknowledgments", level=1)
    doc.add_paragraph(_strip_html(text))
    doc.add_page_break()


def _make_table_of_contents_page(doc, toc_data):
    """
    If you want a fully auto-generated TOC, python-docx doesn't do it well.
    We'll do a placeholder approach.
    'toc_data' might be a list of tuples: (heading, page).
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_heading("Table of Contents", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not toc_data:
        doc.add_paragraph("No table of contents provided.")
    else:
        for item in toc_data:
            heading_str = item.get("heading", "Untitled Heading")
            page_str = item.get("page", "??")
            doc.add_paragraph(f"{heading_str} ...... {page_str}")
    doc.add_page_break()


def _make_list_of_tables_page(doc, lot_data):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_heading("List of Tables", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not lot_data:
        doc.add_paragraph("No tables listed.")
    else:
        for item in lot_data:
            label = item.get("label", "Table ?")
            page_str = item.get("page", "??")
            doc.add_paragraph(f"{label} ...... {page_str}")
    doc.add_page_break()


def _make_list_of_figures_page(doc, lof_data):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_heading("List of Figures", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not lof_data:
        doc.add_paragraph("No figures listed.")
    else:
        for item in lof_data:
            label = item.get("label", "Figure ?")
            page_str = item.get("page", "??")
            doc.add_paragraph(f"{label} ...... {page_str}")
    doc.add_page_break()


# Simple HTML stripper
def _strip_html(text):
    return re.sub(r"<[^>]*>", "", text or "").strip()
