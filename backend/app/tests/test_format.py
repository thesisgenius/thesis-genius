from unittest import mock

import pytest
from app.models.data import Thesis, User
from app.services.apaservice import generate_apa_formatted_document
from app.utils.formatter import format_to_apa


@pytest.fixture
def thesis_data():
    """
    Fixture to provide a sample thesis data structure for testing.
    """
    return {
        "title": "Sample Thesis Title",
        "abstract": "This is a sample abstract for the thesis.",
        "content": "This is the main content of the thesis.",
        "student": User(
            id=1,
            first_name="John",
            last_name="Doe",
            username="johndoe",
            email="john@example.com",
            password="hashed_password",
            role_id=1,
        ),
        "references": [
            mock.Mock(
                author="Jane Doe",
                title="Sample Reference Title",
                journal="Journal of Testing",
                publication_year=2023,
                publisher="Testing Press",
                doi="10.1234/testing.5678",
            )
        ],
    }


@mock.patch("app.services.apaservice.Thesis")
def test_generate_apa_formatted_document(mock_thesis, thesis_data):
    """
    Test the service function that generates an APA-formatted document.
    """
    # Mock reference objects
    mock_references = [
        mock.Mock(
            author="Jane Doe",
            title="Sample Reference Title",
            journal="Journal of Testing",
            publication_year=2023,
            publisher="Testing Press",
            doi="10.1234/testing.5678",
        )
    ]

    # Mock thesis and references
    mock_thesis.get_or_none.return_value = mock.Mock(
        title=thesis_data["title"],
        abstract=thesis_data["abstract"],
        content=thesis_data["content"],
        student=mock.Mock(
            username="john.doe@example.com", first_name="John", last_name="Doe"
        ),
        references=mock_references,
    )

    # Call the service
    doc_path = generate_apa_formatted_document(1)

    # Assertions
    assert doc_path == "/tmp/apa_formatted_thesis.docx"
    mock_thesis.get_or_none.assert_called_once_with(Thesis.id == 1)


def test_format_to_apa(thesis_data):
    """
    Test the utility function that formats a document into APA style.
    """
    # Extract thesis metadata and references
    thesis_metadata = {
        "title": thesis_data["title"],
        "student": thesis_data["student"].first_name
        + " "
        + thesis_data["student"].last_name,
        "institution": "ThesisGenius University",
        "abstract": thesis_data["abstract"],
        "content": thesis_data["content"],
    }
    references = [
        {
            "author": ref.author,
            "title": ref.title,
            "journal": ref.journal,
            "publication_year": ref.publication_year,
            "publisher": ref.publisher,
            "doi": ref.doi,
        }
        for ref in thesis_data["references"]
    ]

    # Call the utility function
    doc_path = format_to_apa(thesis_metadata, references)

    # Assertions
    assert doc_path == "/tmp/apa_formatted_thesis.docx"

    # Verify file content
    from docx import Document

    doc = Document(doc_path)
    assert thesis_metadata["title"] in [p.text for p in doc.paragraphs]
    assert thesis_metadata["abstract"] in [p.text for p in doc.paragraphs]
    assert "References" in [
        h.text for h in doc.paragraphs if h.style.name == "Heading 1"
    ]
    assert any(ref["author"] in p.text for p in doc.paragraphs for ref in references)


@mock.patch("app.routes.format.generate_apa_formatted_document")
def test_format_to_apa_endpoint(mock_generate_doc, client):
    """
    Test the APA formatter endpoint.
    """
    # Mock the service response
    mock_generate_doc.return_value = "/tmp/apa_formatted_thesis.docx"

    # Call the endpoint
    response = client.get("/api/format/apa/1")

    # Assertions
    assert response.status_code == 200
    assert response.headers["Content-Disposition"] in [
        'attachment; filename="APA_Thesis.docx"',
        "attachment; filename=APA_Thesis.docx",
    ]
    assert (
        response.content_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    mock_generate_doc.assert_called_once_with(1)


def test_format_to_apa_endpoint_thesis_not_found(client):
    """
    Test the APA formatter endpoint when the thesis is not found.
    """
    response = client.get("/api/format/apa/99999")

    # Assertions
    assert response.status_code == 404
    assert response.json["error"] == "Thesis with ID 99999 not found."


@mock.patch("app.routes.format.generate_apa_formatted_document")
def test_format_to_apa_endpoint_server_error(mock_generate_doc, client):
    """
    Test the APA formatter endpoint when an exception occurs.
    """
    # Mock the service to raise an exception
    mock_generate_doc.side_effect = Exception("Unexpected error")

    # Call the endpoint
    response = client.get("/api/format/apa/1")

    # Assertions
    assert response.status_code == 500
    assert "An error occurred" in response.json["error"]
